from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document as DocxDocument
from llama_index.core import Document

from app.config import DATA_DIR
from app.ocr_utils import extract_text_from_pdf


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def iter_documents(data_dir: Path = DATA_DIR) -> Iterable[Path]:
    return (
        path
        for path in sorted(data_dir.rglob("*"))
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )


def read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def read_docx(path: Path) -> str:
    document = DocxDocument(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())


def read_document(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(str(path))
    if suffix == ".docx":
        return read_docx(path)
    if suffix == ".txt":
        return read_txt(path)
    raise ValueError(f"Unsupported file type: {path}")


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" *\n *", "\n", text)
    return text.strip()


def load_documents(data_dir: Path = DATA_DIR) -> list[Document]:
    documents: list[Document] = []
    for path in iter_documents(data_dir):
        text = clean_text(read_document(path))
        if not text:
            continue

        documents.append(
            Document(
                text=text,
                metadata={
                    "filename": path.name,
                    "path": str(path.relative_to(data_dir)),
                    "extension": path.suffix.lower(),
                },
            )
        )
    return documents
